"""Export transcripts to PDF, DOCX, and TXT formats."""

import logging
from dataclasses import dataclass
from pathlib import Path
from typing import List, Optional

from audiosmith.exceptions import DocumentFormattingError
from audiosmith.models import DubbingSegment

logger = logging.getLogger(__name__)


@dataclass
class FormatterOptions:
    """Options for document formatting."""
    title: Optional[str] = None
    include_timestamps: bool = True
    include_speaker_labels: bool = False
    font_size: int = 11
    bilingual: bool = False


class DocumentFormatter:
    """Export transcripts to PDF, DOCX, and TXT."""

    @staticmethod
    def _format_timestamp(start: float, end: float) -> str:
        def fmt(t: float) -> str:
            m, s = int(t // 60), t % 60
            return f"{m}:{s:04.1f}"
        return f"[{fmt(start)} - {fmt(end)}]"

    def _format_segment_text(self, seg: DubbingSegment, options: FormatterOptions) -> str:
        parts: List[str] = []
        if options.include_timestamps:
            parts.append(self._format_timestamp(seg.start_time, seg.end_time))
        if options.include_speaker_labels and seg.speaker_id:
            parts.append(f"[{seg.speaker_id}]")
        parts.append(seg.original_text)
        if options.bilingual and seg.translated_text:
            parts.append(f"({seg.translated_text})")
        return " ".join(parts)

    def to_txt(
        self, segments: List[DubbingSegment], output_path: Path,
        options: Optional[FormatterOptions] = None,
    ) -> Path:
        """Export to UTF-8 plain text."""
        opts = options or FormatterOptions()
        lines: List[str] = []
        if opts.title:
            lines.extend([opts.title, ""])
        for seg in segments:
            lines.extend([self._format_segment_text(seg, opts), ""])
        output_path.write_text("\n".join(lines), encoding="utf-8")
        return output_path

    def to_pdf(
        self, segments: List[DubbingSegment], output_path: Path,
        options: Optional[FormatterOptions] = None,
    ) -> Path:
        """Export to PDF via fpdf2 (lazy import)."""
        opts = options or FormatterOptions()
        try:
            from fpdf import FPDF
        except ImportError:
            raise DocumentFormattingError(
                "fpdf2 required for PDF export. Install: pip install fpdf2"
            )
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        if opts.title:
            pdf.set_font("Helvetica", "B", 16)
            pdf.cell(0, 10, opts.title, new_x="LMARGIN", new_y="NEXT", align="C")
            pdf.ln(5)
        pdf.set_font("Helvetica", size=opts.font_size)
        for seg in segments:
            pdf.multi_cell(0, 5, self._format_segment_text(seg, opts))
            pdf.ln(2)
        pdf.output(str(output_path))
        return output_path

    def to_docx(
        self, segments: List[DubbingSegment], output_path: Path,
        options: Optional[FormatterOptions] = None,
    ) -> Path:
        """Export to DOCX via python-docx (lazy import)."""
        opts = options or FormatterOptions()
        try:
            from docx import Document
            from docx.shared import Pt
        except ImportError:
            raise DocumentFormattingError(
                "python-docx required for DOCX export. Install: pip install python-docx"
            )
        doc = Document()
        if opts.title:
            doc.add_heading(opts.title, 0)
        for seg in segments:
            para = doc.add_paragraph(self._format_segment_text(seg, opts))
            if opts.font_size != 11:
                for run in para.runs:
                    run.font.size = Pt(opts.font_size)
        doc.save(str(output_path))
        return output_path
